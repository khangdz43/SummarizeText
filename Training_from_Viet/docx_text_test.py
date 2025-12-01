from docx import Document

# Tạo document mới
doc = Document()

# Thêm title
doc.add_heading("Văn bản thử nghiệm tóm tắt", level=1)

# Nội dung mẫu: 6 đoạn
paragraphs = [
    "Đoạn 1: Đây là một văn bản mẫu dùng để thử nghiệm chức năng tóm tắt văn bản. "
    "Bạn có thể thêm nhiều đoạn khác nhau để mô phỏng văn bản thật.",
    
    "Đoạn 2: Flask sẽ đọc file này và lấy nội dung để gửi vào mô-đun tóm tắt. "
    "Mỗi đoạn nên có vài câu để kiểm tra khả năng tóm tắt theo tỷ lệ.",
    
    "Đoạn 3: Việc tóm tắt văn bản giúp rút gọn thông tin quan trọng mà vẫn giữ ý chính. "
    "Đây là chức năng hữu ích trong nhiều ứng dụng học tập và nghiên cứu.",
    
    "Đoạn 4: Bạn có thể thử thay đổi nội dung để xem kết quả tóm tắt thay đổi như thế nào. "
    "Các thuật toán tóm tắt có thể khác nhau về độ dài và ngữ cảnh.",
    
    "Đoạn 5: Việc thử nghiệm với file Word giúp kiểm tra tính ổn định của ứng dụng Flask "
    "trước khi triển khai thực tế.",
    
    "Đoạn 6: Đây là đoạn cuối cùng. Nội dung này đảm bảo văn bản đủ dài để thử nghiệm "
    "các tỷ lệ tóm tắt khác nhau, từ 20% đến 50%."
]

# Thêm các đoạn vào document
for para in paragraphs:
    doc.add_paragraph(para)

# Lưu file
filename = "test_summary_full.docx"
doc.save(filename)

print(f"File '{filename}' đã được tạo xong. Bạn có thể dùng để thử nghiệm tóm tắt.")
