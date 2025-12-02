import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(__file__)))

from flask import Flask, render_template, request, jsonify
import summarizeController as summarizeController
from docx import Document
import requests
from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = "uploads"
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route("/")
def home():
    return render_template("index.html", summary_result=None)

@app.route("/submit", methods=["POST"])
def submit():
    scale = int(request.form.get("scale", 30))
    contents = request.form.get("contents", "").strip()

    # 1. Xử lý file Word upload
    word_file = request.files.get("file")
    if word_file and word_file.filename != "" and word_file.filename.endswith(".docx"):
        filename = secure_filename(word_file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        word_file.save(filepath)
        doc = Document(filepath)
        file_text = "\n".join([p.text for p in doc.paragraphs])
        contents = file_text # ưu tiên nội dung từ file Word

    
    # 2. Xử lý URL nếu có
    if not contents:  # chỉ lấy URL nếu textarea trống và file không có
        url = request.form.get("url", "").strip()
        if url:
            try:
                html = requests.get(url).text
                soup = BeautifulSoup(html, "html.parser")
                for tag in soup(["script", "style"]):
                    tag.extract()
                text = soup.get_text(separator="\n")
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                contents = "\n".join(lines)
            except Exception as e:
                return render_template("index.html", summary_result=f"Không thể lấy nội dung từ URL: {str(e)}")

    # 3. Nếu không có nội dung nào
    if not contents:
        return render_template("index.html", summary_result="Không có nội dung để tóm tắt!")

    # 4. Tóm tắt
    summarization = summarizeController.summarization(contents, scale)
    return render_template("index.html", summary_result=summarization)

# Phần của API JSON, file docx 
@app.route("/submit_json", methods=["POST"])
def submit_json():
    scale = int(request.form.get("scale", 30))
    contents = request.form.get("contents", "").strip()

    # 1. Xử lý file Word upload
    word_file = request.files.get("file")
    if word_file and word_file.filename != "" and word_file.filename.endswith(".docx"):
        from werkzeug.utils import secure_filename
        from docx import Document
        filename = secure_filename(word_file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        word_file.save(filepath)
        doc = Document(filepath)
        file_text = "\n".join([p.text for p in doc.paragraphs])
        contents = file_text  # ưu tiên nội dung từ file Word

    # 2. Xử lý URL nếu textarea rỗng và không có file
    if not contents:
        url = request.form.get("url", "").strip()
        if url:
            import requests
            from bs4 import BeautifulSoup
            try:
                html = requests.get(url).text
                soup = BeautifulSoup(html, "html.parser")
                for tag in soup(["script", "style"]):
                    tag.extract()
                text = soup.get_text(separator="\n")
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                contents = "\n".join(lines)
            except Exception as e:
                return jsonify({"success": False, "error": f"Không thể lấy nội dung từ URL: {str(e)}"})

    # 3. Nếu không có nội dung
    if not contents:
        return jsonify({"success": False, "error": "Không có nội dung để tóm tắt!"})

    # 4. Tóm tắt
    import summarizeController
    summarization = summarizeController.summarization(contents, scale)
    return jsonify({"success": True, "summary": summarization})

if __name__ == "__main__":
    app.run(debug=True)
